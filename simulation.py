import os
import io
import sys
import cv2
import gym
import base64
import asyncio
import panda_gym
import numpy as np
from tqdm import tqdm
from PIL import Image
from aiohttp import web
from docx import Document
from docx.shared import Inches
from datetime import datetime



from robot import Robot
from core import AbstractSimulation, BASE_DIR
from config.config import SimulationConfig, RobotConfig


class Simulation(AbstractSimulation):
    def __init__(self, cfg=SimulationConfig()) -> None:
        #super().__init__(cfg)

        self.cfg = cfg
        # init env
        self.env = gym.make(f"Panda{cfg.env_name}-v2", render=cfg.render, debug=cfg.debug)
        # init robots
        # count number of tasks solved from a plan 
        self.task_counter = 0
        self.prev_OD_response = ""

        # simulation time
        self.t = 0.
        env_info = (self.env.robots_info, self.env.objects_info)
        self.robot = Robot(env_info,RobotConfig(self.cfg.task))
        # count number of tasks solved from a plan 
        self.task_counter = 0
        # bool for stopping simulation
        self.stop_thread = False
        # whether to save frame (initialized to false)
        self.save_video = False
        # init list of RGB frames if wanna save video
        self.frames_list = []
        self.video_name = f"{self.cfg.env_name}_{self.cfg.task}_{datetime.now().strftime('%d-%m-%Y_%H:%M:%S')}"
        self.video_path = os.path.join(BASE_DIR, f"videos/{self.video_name}.mp4")
        # init log file
        if self.cfg.logging:
            self.doc_path = os.path.join(BASE_DIR, f"logs/{self.video_name}.docx")
            self.doc = Document()
            self.doc.save(self.doc_path)

    def _round_list(self, l, n=2):
        """ round list and if the result is -0.0 convert it to 0.0 """
        return [round(x, n) if round(x, n) != -0.0 else 0.0 for x in l]
    
    def _create_scene_description(self):
        """ Look at the observation and create a string that describes the scene to be passed to the task planner """
        ri = 0
        description = "The following is the description of the current scene. Only use this description to decide what to do next.\n"
        for name in self.observation.keys():
            if name.startswith("robot"):
                description += f"- The center of gripper (between the 2 fingers) of the {name} is located at {self._round_list(self.observation[name][:3])}.\n"
                description += f"- The gripper fingers are {round(self.env.robots[ri].get_fingers_width(),2)} apart and the last instruction given to the gripper is {'open_gripper()' if self.robot.gripper>-1 else 'close_gripper()'}.\n"
                ri += 1
            elif name.endswith("_cube"):
                description += f"- The center of the {name[:-5]} cube is located at {self._round_list(self.observation[name])}\n"
            else:
                pass

        return description

        
                
    def _plan_task(self, user_message:str) -> str:
        # retrieve current frame
        """
        frame_np = np.array(self.env.render("rgb_array", 
                                            width=self.cfg.frame_width, height=self.cfg.frame_height,
                                            target_position=self.cfg.frame_target_position,
                                            distance=self.cfg.frame_distance,
                                            yaw=self.cfg.frame_yaw,
                                            pitch=self.cfg.frame_pitch))
        frame_np = frame_np.reshape(self.cfg.frame_width, self.cfg.frame_height, 4).astype(np.uint8)
        frame_np = cv2.cvtColor(frame_np, cv2.COLOR_RGBA2BGR)
        # convert to base64
        _, buffer = cv2.imencode('.jpg', frame_np)
        frame = base64.b64encode(buffer).decode('utf-8')
        """
        scene_desctiption = self._create_scene_description()
        # run VLM
        task:str = self.robot.plan_task(f"{scene_desctiption}\n\n {user_message}")

        if self.cfg.logging:
            self._add_text_to_doc("SYSTEM: " + f"{scene_desctiption}\n\n {user_message}\nPrevious robot action was: {self.prev_OD_response}")
            #self._add_image_to_doc(frame_np)
            self._add_text_to_doc("TP: " + task)
        return task
    
    def _solve_task(self, task:str):
        AI_response = self.robot.solve_task(task)
        if self.cfg.logging:
            self._add_text_to_doc("OD: " + AI_response)

        self.prev_OD_response = AI_response
        return AI_response

    def reset(self):
        # reset pand env
        self.observation = self.env.reset()
        # reset controller
        self.robot.init_states(self.observation, self.t)
        # count number of tasks solved from a plan 
        self.task_counter = 0
        # init list of RGB frames if wanna save video
        self.frames_list = []

    def step(self):
        # increase timestep
        self.t += self.cfg.dt
        # update controller (i.e. set the current gripper position)
        self.robot.init_states(self.observation, self.t)
        # compute action
        action = self.robot.step() # TODO: this is a list because the env may have multiple robots
        if self.cfg.debug:
            trajectory = self.robot.retrieve_trajectory()
            self.env.visualize_trajectory(trajectory)
        # apply action
        self.observation, _, done, _ = self.env.step(action)
        # store RGB frames if wanna save video
        if self.save_video:
            frame = np.array(self.env.render("rgb_array", width=self.cfg.width, height=self.cfg.height))
            frame = frame.reshape(self.cfg.width, self.cfg.height, 4).astype(np.uint8)
            self.frames_list.append(frame)

        return done

    def close(self):
        # close environment
        #self.thread.join()
        self.stop_thread = True
        self.thread.join()
        # init list of RGB frames if wanna save video
        if self.save_video:
            self._save_video()
        # exit
        sys.exit()  

    def _save_video(self):
        # Define the parameters
        fourcc = cv2.VideoWriter_fourcc(*'mp4v')
        # Create a VideoWriter object
        out = cv2.VideoWriter(self.video_path, fourcc, self.cfg.fps, (self.cfg.width, self.cfg.height))
        # Write frames to the video
        for frame in tqdm(self.frames_list):
            # Ensure the frame is in the correct format (RGBA)
            if frame.shape[2] == 3:
                frame = cv2.cvtColor(frame, cv2.COLOR_RGB2RGBA)
            # Convert the frame to BGR format (required by VideoWriter)
            frame_bgr = cv2.cvtColor(frame, cv2.COLOR_RGBA2BGR)
            out.write(frame_bgr)
        # Release the VideoWriter
        out.release()

    def _add_text_to_doc(self, text):
        self.doc.add_paragraph(text)
        self.doc.save(self.doc_path)

    def _add_image_to_doc(self, image):
        image = Image.fromarray(image)
        # Save the PIL image to a bytes buffer
        image_buffer = io.BytesIO()
        image.save(image_buffer, format="PNG")  # You can use JPEG or PNG
        image_buffer.seek(0)  # Move to the beginning of the buffer

        # Add the image from the bytes buffer to the document
        self.doc.add_picture(image_buffer, width=Inches(4))  # Adjust width as necessary
        
        self.doc.save(self.doc_path)

    async def http_plan_task(self, request):
        data = await request.json()
        user_message = data.get('content')
        task = self._plan_task(user_message)
        return web.json_response({"avatar": "TP", "content": task})

    async def http_solve_task(self, request):
        data = await request.json()
        user_task = data.get('content')
        AI_response = self._solve_task(user_task)
        return web.json_response({"avatar": "OD", "content": AI_response})
    
    async def http_save_recording(self, request):
        self.save_video = False
        self._save_video()
        return web.json_response({"response": "Recording saved"})
    
    async def http_start_recording(self, request):      
        self.save_video = True
        return web.json_response({"response": "Recording started"})
    
    async def http_cancel_recording(self, request):
        self.save_video = False
        self.frames_list = []
        return web.json_response({"response": "Recording cancelled"})
    
    async def main(self, app):
        runner = web.AppRunner(app)
        await runner.setup()
        site = web.TCPSite(runner, 'localhost', 8080)
        await site.start()

        await self._run()

    async def _run(self):
        self.reset()
        while not self.stop_thread:
            done = self.step()
            await asyncio.sleep(0.05)
            if done:
                break
        self.env.close()

    def run(self):
        app = web.Application()
        app.add_routes([
            web.post('/plan_task', self.http_plan_task),
            web.post('/solve_task', self.http_solve_task),
            web.get('/save_recording', self.http_save_recording),
            web.get('/start_recording', self.http_start_recording),
            web.get('/cancel_recording', self.http_cancel_recording)
        ])

        asyncio.run(self.main(app))
  

if __name__=="__main__":
    s = Simulation()
    s.run()